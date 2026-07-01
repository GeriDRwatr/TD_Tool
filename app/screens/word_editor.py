import logging
import os

from PySide6 import QtCore, QtGui, QtWidgets

from ..theme import THEME_MGR

try:
    from docx import Document as DocxDocument

    from ..docx_convert import docx_to_html
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

_log = logging.getLogger(__name__)


class WordEditor(QtWidgets.QWidget):
    """Embedded Word document editor — workspace screen."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._path: str | None = None
        self._modified = False
        self._build_ui()
        self.apply_theme()

    # ── public API ────────────────────────────────────────────────────────────

    def open_file(self, path: str):
        self._load(path)

    def has_unsaved_changes(self) -> bool:
        return self._modified

    def has_file(self) -> bool:
        return self._path is not None

    # ── UI ────────────────────────────────────────────────────────────────────

    def _build_ui(self):
        root = QtWidgets.QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        self._toolbar = QtWidgets.QWidget()
        self._toolbar.setFixedHeight(46)
        tbl = QtWidgets.QHBoxLayout(self._toolbar)
        tbl.setContentsMargins(14, 0, 14, 0)
        tbl.setSpacing(6)

        self._btn_open  = self._make_btn("Відкрити...")
        self._btn_save  = self._make_btn("Зберегти")
        self._btn_print = self._make_btn("Друк")
        self._btn_pdf   = self._make_btn("→ PDF")

        tbl.addWidget(self._btn_open)
        tbl.addSpacing(6)
        tbl.addWidget(self._btn_save)
        tbl.addWidget(self._btn_print)
        tbl.addWidget(self._btn_pdf)
        tbl.addStretch()

        self._lbl_title = QtWidgets.QLabel("Немає документа")
        self._lbl_title.setStyleSheet(
            "color: rgba(255,255,255,0.5); font-size: 12px;"
            " background: transparent; border: none;"
        )
        tbl.addWidget(self._lbl_title)
        root.addWidget(self._toolbar)

        self._editor = QtWidgets.QTextEdit()
        self._editor.setAcceptRichText(True)
        self._editor.document().contentsChanged.connect(self._mark_modified)
        root.addWidget(self._editor, 1)

        self._btn_open.clicked.connect(self._on_open)
        self._btn_save.clicked.connect(self._on_save)
        self._btn_print.clicked.connect(self._on_print)
        self._btn_pdf.clicked.connect(self._on_to_pdf)

    def _make_btn(self, text: str) -> QtWidgets.QPushButton:
        btn = QtWidgets.QPushButton(text)
        btn.setFixedHeight(30)
        btn.setCursor(QtCore.Qt.PointingHandCursor)
        return btn

    # ── theme / appearance ────────────────────────────────────────────────────

    def apply_theme(self):
        t = THEME_MGR.get()
        self.setStyleSheet(f"background: {t.viewer_bg};")
        self._toolbar.setStyleSheet(
            f"QWidget {{ background: {t.bg_sidebar};"
            f" border-bottom: 1px solid {t.bg_border}; }}"
        )
        btn_ss = (
            "QPushButton { background: rgba(255,255,255,0.08);"
            " color: rgba(255,255,255,0.85); border: none; border-radius: 5px;"
            " padding: 0 14px; font-size: 13px; }"
            " QPushButton:hover { background: rgba(255,255,255,0.15); }"
            " QPushButton:pressed { background: rgba(255,255,255,0.05); }"
        )
        for btn in [self._btn_open, self._btn_save, self._btn_print, self._btn_pdf]:
            btn.setStyleSheet(btn_ss)

        # Gray canvas, dark scrollbar on canvas
        self._editor.setStyleSheet("""
            QTextEdit {
                background: #c8c8c8;
                border: none;
            }
            QScrollBar:vertical {
                width: 8px; background: transparent; border: none; margin: 0;
            }
            QScrollBar::handle:vertical {
                background: rgba(0,0,0,0.28);
                border-radius: 4px; min-height: 28px;
            }
            QScrollBar::handle:vertical:hover {
                background: rgba(0,0,0,0.45);
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: none; }
            QScrollBar:horizontal {
                height: 8px; background: transparent; border: none; margin: 0;
            }
            QScrollBar::handle:horizontal {
                background: rgba(0,0,0,0.28);
                border-radius: 4px; min-width: 28px;
            }
            QScrollBar::handle:horizontal:hover { background: rgba(0,0,0,0.45); }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0; }
            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal { background: none; }
        """)
        self._apply_page_format()

    def _apply_page_format(self):
        """White page on gray canvas via root frame format."""
        fmt = QtGui.QTextFrameFormat()
        fmt.setBackground(QtGui.QColor("#ffffff"))
        fmt.setMargin(20)   # gray canvas border around the white page
        fmt.setPadding(54)  # page margins inside the white area
        self._editor.document().rootFrame().setFrameFormat(fmt)

    # ── load ──────────────────────────────────────────────────────────────────

    def _load(self, path: str):
        if not HAS_DOCX:
            self._no_docx_error()
            return
        if path.lower().endswith(".doc") and not path.lower().endswith(".docx"):
            QtWidgets.QMessageBox.warning(
                self, "Непідтримуваний формат",
                "Старий формат .doc не підтримується.\n"
                "Відкрийте файл у Microsoft Word або LibreOffice і збережіть як .docx, "
                "потім спробуйте знову."
            )
            return
        try:
            doc = DocxDocument(path)
            html = docx_to_html(doc)
            self._editor.setHtml(html)
            self._apply_page_format()   # restore after setHtml resets frame
            self._editor.moveCursor(QtGui.QTextCursor.Start)
            self._path = path
            self._modified = False
            self._lbl_title.setText(os.path.basename(path))
        except Exception as e:
            QtWidgets.QMessageBox.critical(
                self, "Помилка", f"Не вдалося відкрити файл:\n{e}"
            )

    # ── save ──────────────────────────────────────────────────────────────────

    def _on_open(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self, "Відкрити Word документ", "", "Word files (*.docx)"
        )
        if path:
            self._load(path)

    def _on_save(self):
        if not self._path:
            path, _ = QtWidgets.QFileDialog.getSaveFileName(
                self, "Зберегти Word документ", "", "Word files (*.docx)"
            )
            if not path:
                return
            if not path.lower().endswith(".docx"):
                path += ".docx"
            self._path = path
        self._save_to(self._path)

    def _save_to(self, path: str):
        if not HAS_DOCX:
            self._no_docx_error()
            return
        try:
            doc = DocxDocument()
            qt_doc = self._editor.document()
            block = qt_doc.begin()
            while block.isValid():
                heading = block.blockFormat().headingLevel()
                para = (doc.add_heading(level=min(heading, 9))
                        if heading > 0 else doc.add_paragraph())
                # Alignment
                try:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    a = block.blockFormat().alignment()
                    if a == QtCore.Qt.AlignLeft:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif a == QtCore.Qt.AlignRight:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif a == QtCore.Qt.AlignHCenter:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif a == QtCore.Qt.AlignJustify:
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                except Exception:
                    _log.debug("Не вдалося застосувати вирівнювання абзацу", exc_info=True)
                it = block.begin()
                while not it.atEnd():
                    frag = it.fragment()
                    if frag.isValid():
                        fmt = frag.charFormat()
                        run = para.add_run(frag.text())
                        run.bold      = fmt.fontWeight() >= 700
                        run.italic    = fmt.fontItalic()
                        run.underline = fmt.fontUnderline()
                        if fmt.fontFamily():
                            run.font.name = fmt.fontFamily()
                        if fmt.fontPointSize() > 0:
                            from docx.shared import Pt
                            run.font.size = Pt(fmt.fontPointSize())
                        c = fmt.foreground().color()
                        if c.isValid() and c.alpha() > 0:
                            from docx.shared import RGBColor
                            try:
                                run.font.color.rgb = RGBColor(c.red(), c.green(), c.blue())
                            except Exception:
                                _log.debug("Не вдалося застосувати колір тексту", exc_info=True)
                    it += 1
                block = block.next()
            doc.save(path)
            self._modified = False
            self._lbl_title.setText(os.path.basename(path))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Помилка збереження", f"{e}")

    def _mark_modified(self):
        if self._path is not None and not self._modified:
            self._modified = True

    # ── print ─────────────────────────────────────────────────────────────────

    def _on_print(self):
        from PySide6.QtPrintSupport import QPrintDialog, QPrinter
        printer = QPrinter(QPrinter.HighResolution)
        if QPrintDialog(printer, self).exec() == QtWidgets.QDialog.Accepted:
            self._editor.print_(printer)

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _on_to_pdf(self):
        if self._modified or not self._path:
            if QtWidgets.QMessageBox.question(
                self, "Зберегти?",
                "Для конвертування документ потрібно зберегти. Зберегти зараз?",
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No,
            ) != QtWidgets.QMessageBox.Yes:
                return
            self._on_save()
            if not self._path:
                return

        default = os.path.splitext(self._path)[0] + ".pdf"
        out, _ = QtWidgets.QFileDialog.getSaveFileName(
            self, "Зберегти як PDF", default, "PDF files (*.pdf)"
        )
        if not out:
            return

        src = self._path
        try:
            if self._convert_via_libreoffice(src, out) or self._convert_via_docx2pdf(src, out):
                pass
            else:
                self._qt_print_to_pdf(out)
            QtWidgets.QMessageBox.information(self, "PDF", f"Збережено:\n{out}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Помилка конвертування", f"{e}")

    def _convert_via_libreoffice(self, src: str, out: str) -> bool:
        import subprocess
        try:
            r = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", os.path.dirname(os.path.abspath(out)), src],
                capture_output=True, timeout=60,
            )
            if r.returncode != 0:
                return False
            lo = os.path.join(
                os.path.dirname(os.path.abspath(out)),
                os.path.splitext(os.path.basename(src))[0] + ".pdf",
            )
            if os.path.abspath(lo) != os.path.abspath(out) and os.path.exists(lo):
                os.replace(lo, out)
            return True
        except (FileNotFoundError, subprocess.TimeoutExpired):
            return False

    def _convert_via_docx2pdf(self, src: str, out: str) -> bool:
        try:
            import docx2pdf
            docx2pdf.convert(src, out)
            return True
        except ImportError:
            return False

    def _qt_print_to_pdf(self, path: str):
        from PySide6.QtPrintSupport import QPrinter
        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(path)
        self._editor.print_(printer)
        QtWidgets.QMessageBox.information(self, "PDF", f"Збережено:\n{path}")

    # ── helpers ───────────────────────────────────────────────────────────────

    def _no_docx_error(self):
        QtWidgets.QMessageBox.critical(
            self, "Відсутня залежність",
            "Бібліотека python-docx не встановлена.\n\npip install python-docx",
        )
