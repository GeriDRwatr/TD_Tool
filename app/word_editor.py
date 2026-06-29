import os
import sys
from PySide6 import QtWidgets, QtCore, QtGui
from .theme import THEME_MGR

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class WordEditor(QtWidgets.QWidget):
    """Embedded Word document editor — workspace screen."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._path: str | None = None
        self._modified = False
        self._build_ui()
        self.apply_theme()

    # ── public API ────────────────────────────────────────────────────────────

    def open_file(self, path: str):
        self._load(path)

    def has_unsaved_changes(self) -> bool:
        return self._modified

    def has_file(self) -> bool:
        return self._path is not None

    # ── UI ────────────────────────────────────────────────────────────────────

    def _build_ui(self):
        root = QtWidgets.QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        self._toolbar = QtWidgets.QWidget()
        self._toolbar.setFixedHeight(46)
        tbl = QtWidgets.QHBoxLayout(self._toolbar)
        tbl.setContentsMargins(14, 0, 14, 0)
        tbl.setSpacing(6)

        self._btn_open  = self._make_btn("Відкрити...")
        self._btn_save  = self._make_btn("Зберегти")
        self._btn_print = self._make_btn("Друк")
        self._btn_pdf   = self._make_btn("→ PDF")

        tbl.addWidget(self._btn_open)
        tbl.addSpacing(6)
        tbl.addWidget(self._btn_save)
        tbl.addWidget(self._btn_print)
        tbl.addWidget(self._btn_pdf)
        tbl.addStretch()

        self._lbl_title = QtWidgets.QLabel("Немає документа")
        self._lbl_title.setStyleSheet(
            "color: rgba(255,255,255,0.5); font-size: 12px;"
            " background: transparent; border: none;"
        )
        tbl.addWidget(self._lbl_title)
        root.addWidget(self._toolbar)

        self._editor = QtWidgets.QTextEdit()
        self._editor.setAcceptRichText(True)
        self._editor.document().contentsChanged.connect(self._mark_modified)
        root.addWidget(self._editor, 1)

        self._btn_open.clicked.connect(self._on_open)
        self._btn_save.clicked.connect(self._on_save)
        self._btn_print.clicked.connect(self._on_print)
        self._btn_pdf.clicked.connect(self._on_to_pdf)

    def _make_btn(self, text: str) -> QtWidgets.QPushButton:
        btn = QtWidgets.QPushButton(text)
        btn.setFixedHeight(30)
        btn.setCursor(QtCore.Qt.PointingHandCursor)
        return btn

    # ── theme / appearance ────────────────────────────────────────────────────

    def apply_theme(self):
        t = THEME_MGR.get()
        self.setStyleSheet(f"background: {t.viewer_bg};")
        self._toolbar.setStyleSheet(
            f"QWidget {{ background: {t.bg_sidebar};"
            f" border-bottom: 1px solid {t.bg_border}; }}"
        )
        btn_ss = (
            "QPushButton { background: rgba(255,255,255,0.08);"
            " color: rgba(255,255,255,0.85); border: none; border-radius: 5px;"
            " padding: 0 14px; font-size: 13px; }"
            " QPushButton:hover { background: rgba(255,255,255,0.15); }"
            " QPushButton:pressed { background: rgba(255,255,255,0.05); }"
        )
        for btn in [self._btn_open, self._btn_save, self._btn_print, self._btn_pdf]:
            btn.setStyleSheet(btn_ss)

        # Gray canvas, dark scrollbar on canvas
        self._editor.setStyleSheet("""
            QTextEdit {
                background: #c8c8c8;
                border: none;
            }
            QScrollBar:vertical {
                width: 8px; background: transparent; border: none; margin: 0;
            }
            QScrollBar::handle:vertical {
                background: rgba(0,0,0,0.28);
                border-radius: 4px; min-height: 28px;
            }
            QScrollBar::handle:vertical:hover {
                background: rgba(0,0,0,0.45);
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: none; }
            QScrollBar:horizontal {
                height: 8px; background: transparent; border: none; margin: 0;
            }
            QScrollBar::handle:horizontal {
                background: rgba(0,0,0,0.28);
                border-radius: 4px; min-width: 28px;
            }
            QScrollBar::handle:horizontal:hover { background: rgba(0,0,0,0.45); }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0; }
            QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal { background: none; }
        """)
        self._apply_page_format()

    def _apply_page_format(self):
        """White page on gray canvas via root frame format."""
        fmt = QtGui.QTextFrameFormat()
        fmt.setBackground(QtGui.QColor("#ffffff"))
        fmt.setMargin(20)   # gray canvas border around the white page
        fmt.setPadding(54)  # page margins inside the white area
        self._editor.document().rootFrame().setFrameFormat(fmt)

    # ── load ──────────────────────────────────────────────────────────────────

    def _load(self, path: str):
        if not HAS_DOCX:
            self._no_docx_error()
            return
        try:
            doc = DocxDocument(path)
            html = self._docx_to_html(doc)
            self._editor.setHtml(html)
            self._apply_page_format()   # restore after setHtml resets frame
            self._editor.moveCursor(QtGui.QTextCursor.Start)
            self._path = path
            self._modified = False
            self._lbl_title.setText(os.path.basename(path))
        except Exception as e:
            QtWidgets.QMessageBox.critical(
                self, "Помилка", f"Не вдалося відкрити файл:\n{e}"
            )

    # ── docx → HTML (full formatting) ────────────────────────────────────────

    def _docx_to_html(self, doc) -> str:
        import html as _h
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            ALIGN = {
                WD_ALIGN_PARAGRAPH.LEFT:    'left',
                WD_ALIGN_PARAGRAPH.CENTER:  'center',
                WD_ALIGN_PARAGRAPH.RIGHT:   'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
                None: 'justify',
            }
        except Exception:
            ALIGN = {None: 'justify'}

        list_counters: dict[tuple, int] = {}

        parts = [
            "<html><body style='margin:0; padding:0;"
            " font-family:\"Times New Roman\"; font-size:12pt; color:#111;'>"
        ]

        for para in doc.paragraphs:
            pf = para.paragraph_format
            sname = (para.style.name or '') if para.style else ''
            sl = sname.lower()

            align = ALIGN.get(para.alignment, 'justify')

            # Spacing
            sb = int(pf.space_before.pt) if pf.space_before else 0
            sa = int(pf.space_after.pt)  if pf.space_after  else 6

            # Line spacing
            ls_css = ''
            if pf.line_spacing:
                try:
                    from docx.enum.text import WD_LINE_SPACING
                    if pf.line_spacing_rule in (
                        WD_LINE_SPACING.MULTIPLE,
                        WD_LINE_SPACING.AT_LEAST,
                        None,
                    ):
                        ls_css = f' line-height:{float(pf.line_spacing):.2f};'
                except Exception:
                    pass

            # Indentation (pt)
            li_pt = 0
            fi_pt = 0
            if pf.left_indent:
                try: li_pt = int(pf.left_indent.pt)
                except Exception: pass
            if pf.first_line_indent:
                try: fi_pt = int(pf.first_line_indent.pt)
                except Exception: pass

            # List detection — check para pPr first, then style pPr (inheritance)
            list_prefix = ''
            numPr = None
            pPr = para._p.pPr
            if pPr is not None:
                numPr = pPr.find(f'{_NS}numPr')
            if numPr is None and para.style and para.style.element is not None:
                spPr = para.style.element.find(f'{_NS}pPr')
                if spPr is not None:
                    numPr = spPr.find(f'{_NS}numPr')

            if numPr is not None:
                ilvl_el  = numPr.find(f'{_NS}ilvl')
                numId_el = numPr.find(f'{_NS}numId')
                ilvl  = int(ilvl_el.get(f'{_NS}val',  '0')) if ilvl_el  is not None else 0
                numId = int(numId_el.get(f'{_NS}val', '0')) if numId_el is not None else 0

                if numId > 0:
                    # Reset deeper levels on increment
                    for k in list(list_counters):
                        if k[0] == numId and k[1] > ilvl:
                            list_counters[k] = 0
                    key = (numId, ilvl)
                    list_counters[key] = list_counters.get(key, 0) + 1

                    ltype = self._list_type(doc, numId, ilvl)
                    if ltype == 'decimal':
                        list_prefix = f'{list_counters[key]}.&nbsp;&nbsp;'
                    else:
                        list_prefix = '•&nbsp;&nbsp;'

                    if li_pt == 0:
                        li_pt = 36 + ilvl * 18
                    if fi_pt == 0:
                        fi_pt = -18

            # Build paragraph style
            p_css = (
                f"text-align:{align};"
                f" margin-top:{sb}pt; margin-bottom:{sa}pt;"
                f" margin-left:{li_pt}pt; text-indent:{fi_pt}pt;{ls_css}"
            )

            # Detect heading tag
            if   'heading 1' in sl: tag = 'h1'
            elif 'heading 2' in sl: tag = 'h2'
            elif 'heading 3' in sl: tag = 'h3'
            else:                   tag = 'p'

            # Build runs
            runs_html = [list_prefix] if list_prefix else []
            for run in para.runs:
                text = _h.escape(run.text)
                if not text:
                    continue

                span_css: list[str] = []

                # Font family
                fname = run.font.name
                if fname:
                    span_css.append(f"font-family:'{fname}'")

                # Font size
                fsize = run.font.size
                if fsize:
                    try: span_css.append(f"font-size:{fsize.pt:.1f}pt")
                    except Exception: pass

                # Text color
                try:
                    from docx.enum.dml import MSO_COLOR_TYPE
                    c = run.font.color
                    if c.type == MSO_COLOR_TYPE.RGB and c.rgb:
                        span_css.append(f"color:#{c.rgb}")
                except Exception:
                    pass

                if span_css:
                    text = f'<span style="{"; ".join(span_css)}">{text}</span>'
                if run.bold:           text = f'<b>{text}</b>'
                if run.italic:         text = f'<i>{text}</i>'
                if run.underline:      text = f'<u>{text}</u>'
                if run.font.strike:    text = f'<s>{text}</s>'

                runs_html.append(text)

            inner = ''.join(runs_html) or '&nbsp;'
            parts.append(f'<{tag} style="{p_css}">{inner}</{tag}>')

        parts.append('</body></html>')
        return '\n'.join(parts)

    def _list_type(self, doc, numId: int, ilvl: int) -> str:
        """Return 'decimal' or 'bullet' for a list level."""
        try:
            np = doc.part.numbering_part
            if np is None:
                return 'decimal'
            root = np._element
            abstract_id = None
            for num_el in root.findall(f'{_NS}num'):
                if num_el.get(f'{_NS}numId') == str(numId):
                    an = num_el.find(f'{_NS}abstractNumId')
                    if an is not None:
                        abstract_id = an.get(f'{_NS}val')
                    break
            if abstract_id is None:
                return 'decimal'
            for an_el in root.findall(f'{_NS}abstractNum'):
                if an_el.get(f'{_NS}abstractNumId') == abstract_id:
                    for lvl in an_el.findall(f'{_NS}lvl'):
                        if lvl.get(f'{_NS}ilvl') == str(ilvl):
                            fmt_el = lvl.find(f'{_NS}numFmt')
                            if fmt_el is not None:
                                v = fmt_el.get(f'{_NS}val', 'decimal')
                                return 'decimal' if v == 'decimal' else 'bullet'
            return 'decimal'
        except Exception:
            return 'decimal'

    # ── save ──────────────────────────────────────────────────────────────────

    def _on_open(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self, "Відкрити Word документ", "", "Word files (*.docx *.doc)"
        )
        if path:
            self._load(path)

    def _on_save(self):
        if not self._path:
            path, _ = QtWidgets.QFileDialog.getSaveFileName(
                self, "Зберегти Word документ", "", "Word files (*.docx)"
            )
            if not path:
                return
            if not path.lower().endswith(".docx"):
                path += ".docx"
            self._path = path
        self._save_to(self._path)

    def _save_to(self, path: str):
        if not HAS_DOCX:
            self._no_docx_error()
            return
        try:
            doc = DocxDocument()
            qt_doc = self._editor.document()
            block = qt_doc.begin()
            while block.isValid():
                heading = block.blockFormat().headingLevel()
                para = (doc.add_heading(level=min(heading, 9))
                        if heading > 0 else doc.add_paragraph())
                # Alignment
                try:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    a = block.blockFormat().alignment()
                    if a == QtCore.Qt.AlignLeft:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif a == QtCore.Qt.AlignRight:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif a == QtCore.Qt.AlignHCenter:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif a == QtCore.Qt.AlignJustify:
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                except Exception:
                    pass
                it = block.begin()
                while not it.atEnd():
                    frag = it.fragment()
                    if frag.isValid():
                        fmt = frag.charFormat()
                        run = para.add_run(frag.text())
                        run.bold      = fmt.fontWeight() >= 700
                        run.italic    = fmt.fontItalic()
                        run.underline = fmt.fontUnderline()
                        if fmt.fontFamily():
                            run.font.name = fmt.fontFamily()
                        if fmt.fontPointSize() > 0:
                            from docx.shared import Pt
                            run.font.size = Pt(fmt.fontPointSize())
                        c = fmt.foreground().color()
                        if c.isValid() and c.alpha() > 0:
                            from docx.shared import RGBColor
                            try:
                                run.font.color.rgb = RGBColor(c.red(), c.green(), c.blue())
                            except Exception:
                                pass
                    it += 1
                block = block.next()
            doc.save(path)
            self._modified = False
            self._lbl_title.setText(os.path.basename(path))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Помилка збереження", f"{e}")

    def _mark_modified(self):
        if self._path is not None and not self._modified:
            self._modified = True

    # ── print ─────────────────────────────────────────────────────────────────

    def _on_print(self):
        from PySide6.QtPrintSupport import QPrinter, QPrintDialog
        printer = QPrinter(QPrinter.HighResolution)
        if QPrintDialog(printer, self).exec() == QtWidgets.QDialog.Accepted:
            self._editor.print_(printer)

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _on_to_pdf(self):
        if self._modified or not self._path:
            if QtWidgets.QMessageBox.question(
                self, "Зберегти?",
                "Для конвертування документ потрібно зберегти. Зберегти зараз?",
                QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No,
            ) != QtWidgets.QMessageBox.Yes:
                return
            self._on_save()
            if not self._path:
                return

        default = os.path.splitext(self._path)[0] + ".pdf"
        out, _ = QtWidgets.QFileDialog.getSaveFileName(
            self, "Зберегти як PDF", default, "PDF files (*.pdf)"
        )
        if not out:
            return

        try:
            if sys.platform == "win32":
                try:
                    import docx2pdf
                    docx2pdf.convert(self._path, out)
                except ImportError:
                    self._qt_print_to_pdf(out)
                    return
            else:
                import subprocess
                r = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf",
                     "--outdir", os.path.dirname(os.path.abspath(out)), self._path],
                    capture_output=True, timeout=60,
                )
                if r.returncode != 0:
                    raise RuntimeError(r.stderr.decode(errors="replace"))
                lo = os.path.join(
                    os.path.dirname(os.path.abspath(out)),
                    os.path.splitext(os.path.basename(self._path))[0] + ".pdf",
                )
                if os.path.abspath(lo) != os.path.abspath(out) and os.path.exists(lo):
                    os.replace(lo, out)
            QtWidgets.QMessageBox.information(self, "PDF", f"Збережено:\n{out}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Помилка конвертування", f"{e}")

    def _qt_print_to_pdf(self, path: str):
        from PySide6.QtPrintSupport import QPrinter
        printer = QPrinter(QPrinter.HighResolution)
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(path)
        self._editor.print_(printer)
        QtWidgets.QMessageBox.information(self, "PDF", f"Збережено:\n{path}")

    # ── helpers ───────────────────────────────────────────────────────────────

    def _no_docx_error(self):
        QtWidgets.QMessageBox.critical(
            self, "Відсутня залежність",
            "Бібліотека python-docx не встановлена.\n\npip install python-docx",
        )
