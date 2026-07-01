from docx import Document

from app.docx_convert import _heading_tag, _run_to_html, docx_to_html, list_type


def _run(paragraph, text):
    return paragraph.add_run(text)


def test_docx_to_html_wraps_plain_paragraph_in_p_tag():
    doc = Document()
    doc.add_paragraph("Hello world")

    html = docx_to_html(doc)

    assert "<p " in html
    assert "Hello world" in html


def test_docx_to_html_maps_heading_styles_to_tags():
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_heading("Subtitle", level=2)

    html = docx_to_html(doc)

    assert "<h1" in html
    assert "<h2" in html


def test_docx_to_html_escapes_special_characters():
    doc = Document()
    doc.add_paragraph("<script>alert(1)</script> & co")

    html = docx_to_html(doc)

    assert "<script>" not in html
    assert "&lt;script&gt;" in html
    assert "&amp;" in html


def test_docx_to_html_empty_paragraph_renders_nbsp():
    doc = Document()
    doc.add_paragraph("")

    html = docx_to_html(doc)

    assert "&nbsp;" in html


def test_run_to_html_wraps_bold_italic_underline():
    doc = Document()
    p = doc.add_paragraph()
    run = _run(p, "text")
    run.bold = True
    run.italic = True
    run.underline = True

    html = _run_to_html(run)

    assert html == "<u><i><b>text</b></i></u>"


def test_run_to_html_skips_empty_text():
    doc = Document()
    p = doc.add_paragraph()
    run = _run(p, "")

    assert _run_to_html(run) == ""


def test_heading_tag_maps_known_styles_and_falls_back_to_p():
    assert _heading_tag("Heading 1") == "h1"
    assert _heading_tag("Heading 2") == "h2"
    assert _heading_tag("Heading 3") == "h3"
    assert _heading_tag("Normal") == "p"
    assert _heading_tag("") == "p"


def test_list_type_defaults_to_decimal_when_num_id_unknown():
    doc = Document()
    doc.add_paragraph("no lists here")

    assert list_type(doc, num_id=999, ilvl=0) == "decimal"
